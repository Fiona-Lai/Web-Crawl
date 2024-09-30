import requests
from lxml import etree
from selenium import webdriver
import time
from selenium.webdriver.common.by import By
from docx import Document

keyword = input('请输入你想要看的新闻：')
options = webdriver.ChromeOptions()
#去除正在受插件操作
options.add_experimental_option('excludeSwitches',['enable-automation'])
#无痕模式打开
options.add_argument('--incognito')
driver = webdriver.Chrome(executable_path = r'D:\专门写项目\驱动安装\chromedriver.exe',options = options)
driver.get("https://search.globaltimes.cn/QuickSearchCtrl")
#把鼠标定位到输入内容的位置
keyword_ = driver.find_element(By.XPATH,'/html/body/div[3]/div/div/div/div[1]/form/input[2]')
#定位以后输入内容
keyword_.send_keys(keyword)
time.sleep(0.43)
#输入内容以后，还需要点击搜索
click = driver.find_element(By.XPATH,'/html/body/div[3]/div/div/div/div[1]/form/i')
click.click()
time.sleep(3)
driver.maximize_window()
time.sleep(15)
for i in range(4):
    js = "window.scrollBy(0,600)"
    driver.execute_script(js)
    time.sleep(1)
#获取网页源码
data = driver.page_source
html = etree.HTML(data)
link = html.xpath('//div/blockquote/a/@href')
#得到的link是列表形式的xie
time.sleep(10)
#关闭浏览器
driver.quit()
headers = {
    'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/103.0.0.0 Safari/537.36',
    'cookie': 'usprivacy=1---; FXN_flk=1; EID=null; s_ecid=MCMID%7C53025702096023741433609835061423222983; ajs_anonymous_id=3c3d3fca-1034-4fd1-8b38-9d365846e5fd; _gid=GA1.2.1146669240.1657278558; AKA_A2=A; AMCVS_17FC406C5357BA6E0A490D4D%40AdobeOrg=1; AMCV_17FC406C5357BA6E0A490D4D%40AdobeOrg=2121618341%7CMCIDTS%7C19182%7CMCMID%7C53025702096023741433609835061423222983%7CMCAAMLH-1657883283%7C11%7CMCAAMB-1657894322%7CRKhpRz8krg2tLO6pguXWp5olkAcUniQYPHaMWWgdJ3xzPWQmdj0y%7CMCOPTOUT-1657296722s%7CNONE%7CMCAID%7CNONE; s_cc=true; s_pers=%20s_ppn%3Dfnc%253Aroot%253Aroot%253Achannel%7C1657291344199%3B%20omtr_lv%3D1657289549223%7C1751897549223%3B%20omtr_lv_s%3DLess%2520than%25201%2520day%7C1657291349223%3B%20s_nr%3D1657289549224-Repeat%7C1659881549224%3B; s_sq=%5B%5BB%5D%5D; _ga=GA1.1.1721776522.1657278558; _ga_NW9WX3ZPEG=GS1.1.1657289552.2.1.1657289552.60; s_sess=%20s_ppvl%3Dfnc%25253Aroot%25253Aroot%25253Achannel%252C2%252C2%252C577%252C1280%252C577%252C1280%252C720%252C1.5%252CL%3B%20SC_LINKS%3D%3B%20s_ppv%3D%25253A%25253A%25253Aother%252C42%252C42%252C777%252C1280%252C274%252C1280%252C720%252C1.5%252CL%3B; RT="z=1&dm=foxnews.com&si=e8548800-3be9-4bf4-a510-1dc41807f30a&ss=l5cjd6b8&sl=3&tt=w85&bcn=%2F%2F02179918.akstat.io%2F&ld=1lax&ul=56ym"'
}
# with open('hh.text','a',encoding = 'utf-8')as f1:
for i in range(0,10):
    s_response = requests.get(url = link[i],headers=headers)
    s_html = etree.HTML(s_response.content)
    print(link[i])
    title = s_html.xpath("//div/div[@class='article_top']/div[2]/text()")
    print(title)
    article = s_html.xpath("//div/div/p/text()")
    print(article)
    time = s_html.xpath('//div/div/span[2]/text()')
    print(time)
    print("-"*20)
    # file = Document()
    # paragraph3 = file.add_paragraph()
    # file.add_paragraph("标题：")
    # file.add_paragraph(str(title[0]))
    # file.add_paragraph("内容")
    # file.add_paragraph(str(neirong[0]))
    # print(neirong[0])
    # file.add_paragraph("time:")
    # file.add_paragraph(str(time[0]))
    # file.add_paragraph('---------------------------------------------------------------------------')
    # file.save('asdf.docx')