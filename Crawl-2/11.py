from docx import Document

doc = Document()
# 这里相当于输入了一个空格，后面等待着文字输入
paragraph3 = doc.add_paragraph(r"o我哦哦哦哦哦哦哦哦哦")
paragraph3.add_run("我被加粗了文字块儿")
paragraph3 = doc.add_paragraph("啦啦啦啦啦啦啦啦啦啦啦啦啦")
paragraph3.add_run("，我是普通文字块儿，")
paragraph3.add_run("我是斜体文字块儿")
doc.save(r"D:\专门写项目\爬虫练习/学习.docx")